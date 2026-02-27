# -*- coding: utf-8 -*-
"""
读取docx文件内容并保存到文本文件
"""
import sys
from pathlib import Path
from docx import Document

def read_docx_to_file(docx_path, output_path):
    """读取docx文件内容并保存到文本文件"""
    doc = Document(docx_path)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("=" * 60 + "\n")
        f.write(f"文件: {Path(docx_path).name}\n")
        f.write("=" * 60 + "\n\n")

        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:
                f.write(f"{i}. {text}\n")

        f.write("\n" + "=" * 60 + "\n")
        f.write(f"总段落数: {len([p for p in doc.paragraphs if p.text.strip()])}\n")
        f.write("=" * 60 + "\n")

    print(f"内容已保存到: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) > 2:
        docx_path = sys.argv[1]
        output_path = sys.argv[2]
        read_docx_to_file(docx_path, output_path)
    else:
        print("用法: python read_docx_to_file.py <docx文件路径> <输出文件路径>")
