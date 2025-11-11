"""
文档读取器测试脚本

用于测试 Excel, Word, PDF 文件读取功能
"""
import sys
from pathlib import Path

# 添加 backend 到路径
backend_dir = Path(__file__).parent.parent
sys.path.insert(0, str(backend_dir))

from utils.document_reader import DocumentReader, format_document_summary


def test_dependencies():
    """测试依赖是否已安装"""
    print("🔍 检查依赖库安装状态...\n")

    dependencies = {
        "pandas": "Excel 文件支持",
        "openpyxl": "Excel 文件支持",
        "docx": "Word 文档支持",
        "PyPDF2": "PDF 文件支持"
    }

    missing = []
    for package, purpose in dependencies.items():
        try:
            __import__(package)
            print(f"✅ {package:15s} - {purpose}")
        except ImportError:
            print(f"❌ {package:15s} - {purpose} (未安装)")
            missing.append(package)

    if missing:
        print(f"\n⚠️  缺少依赖库: {', '.join(missing)}")
        print("\n安装命令:")
        print(f"pip install {' '.join(missing)}")
        return False
    else:
        print("\n✅ 所有依赖库已安装！")
        return True


def create_test_files():
    """创建测试文件（如果不存在）"""
    test_dir = Path(__file__).parent / "test_files"
    test_dir.mkdir(exist_ok=True)

    print(f"\n📁 测试文件目录: {test_dir}\n")

    # 创建测试 Excel
    try:
        import pandas as pd
        excel_file = test_dir / "test.xlsx"
        if not excel_file.exists():
            df = pd.DataFrame({
                "姓名": ["张三", "李四", "王五"],
                "年龄": [25, 30, 35],
                "城市": ["北京", "上海", "广州"],
                "薪资": [8000, 12000, 15000]
            })
            df.to_excel(excel_file, index=False)
            print(f"✅ 创建测试 Excel: {excel_file}")
        else:
            print(f"✓  测试 Excel 已存在: {excel_file}")
    except ImportError:
        print("⊘  跳过 Excel 测试文件创建（pandas 未安装）")

    # 创建测试 Word
    try:
        from docx import Document
        word_file = test_dir / "test.docx"
        if not word_file.exists():
            doc = Document()
            doc.add_heading('测试文档', 0)
            doc.add_paragraph('这是第一段测试内容。')
            doc.add_paragraph('这是第二段测试内容，包含更多文字。')
            doc.add_heading('子标题', level=1)
            doc.add_paragraph('子标题下的段落。')

            # 添加表格
            table = doc.add_table(rows=3, cols=3)
            table.cell(0, 0).text = "姓名"
            table.cell(0, 1).text = "年龄"
            table.cell(0, 2).text = "城市"
            table.cell(1, 0).text = "张三"
            table.cell(1, 1).text = "25"
            table.cell(1, 2).text = "北京"

            doc.save(word_file)
            print(f"✅ 创建测试 Word: {word_file}")
        else:
            print(f"✓  测试 Word 已存在: {word_file}")
    except ImportError:
        print("⊘  跳过 Word 测试文件创建（python-docx 未安装）")

    print()
    return test_dir


def test_read_excel(test_dir):
    """测试 Excel 读取"""
    excel_file = test_dir / "test.xlsx"
    if not excel_file.exists():
        print("⊘  跳过 Excel 测试（文件不存在）")
        return

    print("=" * 60)
    print("📊 测试 Excel 读取")
    print("=" * 60)

    result = DocumentReader.read_excel(str(excel_file))

    if result.get("success"):
        print(format_document_summary(result))
        print("\n📋 数据预览:")
        for i, row in enumerate(result.get("preview", [])[:3], 1):
            print(f"  第 {i} 行: {row}")
    else:
        print(f"❌ {result.get('error')}")

    print()


def test_read_word(test_dir):
    """测试 Word 读取"""
    word_file = test_dir / "test.docx"
    if not word_file.exists():
        print("⊘  跳过 Word 测试（文件不存在）")
        return

    print("=" * 60)
    print("📝 测试 Word 读取")
    print("=" * 60)

    result = DocumentReader.read_word(str(word_file))

    if result.get("success"):
        print(format_document_summary(result))
        print("\n📋 内容预览:")
        full_text = result.get("full_text", "")
        print(f"  {full_text[:200]}...")
    else:
        print(f"❌ {result.get('error')}")

    print()


def test_read_pdf(test_dir):
    """测试 PDF 读取"""
    # 由于创建 PDF 需要额外的库，这里只测试读取已有的 PDF
    pdf_files = list(test_dir.glob("*.pdf"))

    if not pdf_files:
        print("⊘  跳过 PDF 测试（test_files 目录中没有 PDF 文件）")
        print("   提示：可将任意 PDF 文件放入 test_files 目录进行测试")
        return

    print("=" * 60)
    print("📖 测试 PDF 读取")
    print("=" * 60)

    pdf_file = pdf_files[0]
    result = DocumentReader.read_pdf(str(pdf_file))

    if result.get("success"):
        print(format_document_summary(result))
        print("\n📋 内容预览:")
        full_text = result.get("full_text", "")
        print(f"  {full_text[:200]}...")
    else:
        print(f"❌ {result.get('error')}")

    print()


def test_smart_read(test_dir):
    """测试智能读取（自动识别类型）"""
    print("=" * 60)
    print("🤖 测试智能文档读取")
    print("=" * 60)

    test_files = [
        test_dir / "test.xlsx",
        test_dir / "test.docx"
    ]

    for file_path in test_files:
        if file_path.exists():
            print(f"\n读取: {file_path.name}")
            result = DocumentReader.read_document(str(file_path))
            if result.get("success"):
                print(f"  ✅ 成功 - {result.get('file_name')}")
            else:
                print(f"  ❌ 失败 - {result.get('error')}")

    print()


def main():
    """主测试流程"""
    print("\n" + "=" * 60)
    print("📄 文档读取器测试")
    print("=" * 60 + "\n")

    # 1. 检查依赖
    if not test_dependencies():
        print("\n⚠️  请先安装缺少的依赖库")
        return

    # 2. 创建测试文件
    test_dir = create_test_files()

    # 3. 测试各种文件类型
    test_read_excel(test_dir)
    test_read_word(test_dir)
    test_read_pdf(test_dir)
    test_smart_read(test_dir)

    print("=" * 60)
    print("✅ 测试完成！")
    print("=" * 60)

    print(f"\n💡 提示:")
    print(f"   - 测试文件位于: {test_dir}")
    print(f"   - 可将自己的文件放入该目录进行测试")
    print(f"   - 使用方式: python backend/utils/document_reader.py <文件路径>")


if __name__ == "__main__":
    main()
