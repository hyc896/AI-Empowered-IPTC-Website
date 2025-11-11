"""
文档读取工具模块

支持读取 Excel、Word 和 PDF 文件
"""
from pathlib import Path
from typing import Dict, List, Any, Optional
import logging

logger = logging.getLogger(__name__)


class DocumentReader:
    """通用文档读取器"""

    @staticmethod
    def read_excel(file_path: str, sheet_name: Optional[str] = None) -> Dict[str, Any]:
        """
        读取 Excel 文件

        Args:
            file_path: Excel 文件路径
            sheet_name: 工作表名称（None 表示读取第一个工作表）

        Returns:
            包含数据和元信息的字典
        """
        try:
            import pandas as pd

            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"文件不存在: {file_path}")

            # 读取 Excel
            if sheet_name:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheets = [sheet_name]
            else:
                excel_file = pd.ExcelFile(file_path)
                sheets = excel_file.sheet_names
                df = pd.read_excel(file_path, sheet_name=sheets[0])

            return {
                "success": True,
                "file_name": path.name,
                "file_size": path.stat().st_size,
                "sheets": sheets,
                "current_sheet": sheet_name or sheets[0],
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": list(df.columns),
                "data": df.to_dict('records'),
                "preview": df.head(10).to_dict('records'),
                "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()}
            }

        except ImportError:
            logger.error("pandas 未安装，请运行: pip install pandas openpyxl")
            return {
                "success": False,
                "error": "pandas 库未安装",
                "solution": "pip install pandas openpyxl"
            }
        except Exception as e:
            logger.error(f"读取 Excel 文件失败: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }

    @staticmethod
    def read_word(file_path: str) -> Dict[str, Any]:
        """
        读取 Word 文档

        Args:
            file_path: Word 文件路径

        Returns:
            包含文本和元信息的字典
        """
        try:
            from docx import Document

            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"文件不存在: {file_path}")

            doc = Document(file_path)

            # 提取段落
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name
                    })

            # 提取表格
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    "index": table_idx,
                    "data": table_data
                })

            # 完整文本
            full_text = "\n\n".join([p["text"] for p in paragraphs])

            return {
                "success": True,
                "file_name": path.name,
                "file_size": path.stat().st_size,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "paragraphs": paragraphs,
                "tables": tables,
                "full_text": full_text,
                "text_length": len(full_text)
            }

        except ImportError:
            logger.error("python-docx 未安装，请运行: pip install python-docx")
            return {
                "success": False,
                "error": "python-docx 库未安装",
                "solution": "pip install python-docx"
            }
        except Exception as e:
            logger.error(f"读取 Word 文档失败: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }

    @staticmethod
    def read_pdf(file_path: str) -> Dict[str, Any]:
        """
        读取 PDF 文件

        Args:
            file_path: PDF 文件路径

        Returns:
            包含文本和元信息的字典
        """
        try:
            import PyPDF2

            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"文件不存在: {file_path}")

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                # 提取所有页面文本
                pages = []
                full_text = ""

                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    pages.append({
                        "page_number": page_num + 1,
                        "text": text,
                        "text_length": len(text)
                    })
                    full_text += text + "\n\n"

                # 获取 PDF 元数据
                metadata = pdf_reader.metadata or {}

                return {
                    "success": True,
                    "file_name": path.name,
                    "file_size": path.stat().st_size,
                    "page_count": num_pages,
                    "pages": pages,
                    "full_text": full_text.strip(),
                    "text_length": len(full_text),
                    "metadata": {
                        "title": metadata.get("/Title", ""),
                        "author": metadata.get("/Author", ""),
                        "subject": metadata.get("/Subject", ""),
                        "creator": metadata.get("/Creator", "")
                    }
                }

        except ImportError:
            logger.error("PyPDF2 未安装，请运行: pip install PyPDF2")
            return {
                "success": False,
                "error": "PyPDF2 库未安装",
                "solution": "pip install PyPDF2"
            }
        except Exception as e:
            logger.error(f"读取 PDF 文件失败: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }

    @staticmethod
    def read_document(file_path: str) -> Dict[str, Any]:
        """
        智能读取文档（根据扩展名自动选择方法）

        Args:
            file_path: 文件路径

        Returns:
            包含文件内容和元信息的字典
        """
        path = Path(file_path)

        if not path.exists():
            return {
                "success": False,
                "error": f"文件不存在: {file_path}"
            }

        suffix = path.suffix.lower()

        if suffix in ['.xlsx', '.xls']:
            return DocumentReader.read_excel(file_path)
        elif suffix == '.docx':
            return DocumentReader.read_word(file_path)
        elif suffix == '.pdf':
            return DocumentReader.read_pdf(file_path)
        else:
            return {
                "success": False,
                "error": f"不支持的文件类型: {suffix}",
                "supported_types": [".xlsx", ".xls", ".docx", ".pdf"]
            }


def format_document_summary(result: Dict[str, Any]) -> str:
    """
    格式化文档读取结果为可读的摘要

    Args:
        result: DocumentReader 返回的结果字典

    Returns:
        格式化的文本摘要
    """
    if not result.get("success"):
        return f"❌ 读取失败: {result.get('error')}\n💡 解决方案: {result.get('solution', '请检查文件路径')}"

    file_name = result.get("file_name", "未知")
    file_size = result.get("file_size", 0)
    size_mb = file_size / (1024 * 1024)

    summary = f"📄 文件信息\n"
    summary += f"- 文件名: {file_name}\n"
    summary += f"- 大小: {size_mb:.2f} MB\n"

    # Excel 特定信息
    if "sheets" in result:
        summary += f"- 类型: Excel 工作簿\n"
        summary += f"- 工作表: {', '.join(result['sheets'])}\n"
        summary += f"- 当前表: {result['current_sheet']}\n"
        summary += f"\n📊 数据概览\n"
        summary += f"- 行数: {result['rows']:,}\n"
        summary += f"- 列数: {result['columns']}\n"
        summary += f"- 列名: {', '.join(result['column_names'])}\n"

    # Word 特定信息
    elif "paragraph_count" in result:
        summary += f"- 类型: Word 文档\n"
        summary += f"\n📝 内容概览\n"
        summary += f"- 段落数: {result['paragraph_count']}\n"
        summary += f"- 表格数: {result['table_count']}\n"
        summary += f"- 总字数: {result['text_length']:,}\n"

    # PDF 特定信息
    elif "page_count" in result:
        summary += f"- 类型: PDF 文档\n"
        summary += f"\n📖 内容概览\n"
        summary += f"- 页数: {result['page_count']}\n"
        summary += f"- 总字数: {result['text_length']:,}\n"
        if result.get("metadata", {}).get("title"):
            summary += f"- 标题: {result['metadata']['title']}\n"
        if result.get("metadata", {}).get("author"):
            summary += f"- 作者: {result['metadata']['author']}\n"

    return summary


# 命令行使用示例
if __name__ == "__main__":
    import sys
    import json

    if len(sys.argv) < 2:
        print("用法: python document_reader.py <文件路径>")
        sys.exit(1)

    file_path = sys.argv[1]
    result = DocumentReader.read_document(file_path)

    # 打印摘要
    print(format_document_summary(result))

    # 可选：保存完整结果为 JSON
    if result.get("success"):
        output_file = Path(file_path).stem + "_result.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"\n✅ 完整结果已保存到: {output_file}")
