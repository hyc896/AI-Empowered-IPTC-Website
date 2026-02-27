# -*- coding: utf-8 -*-
"""
从Word文档中提取章节内容并生成JSON格式
用于知识图谱构建
"""
import sys
import json
import re
from pathlib import Path
from docx import Document
from typing import List, Dict, Any


def extract_chapter_structure(doc: Document, book_name: str, book_id: str) -> Dict[str, Any]:
    """
    从Word文档中提取章节结构

    Args:
        doc: Word文档对象
        book_name: 书名
        book_id: 书籍ID

    Returns:
        包含章节内容的字典
    """
    chapters = []
    current_chapter = None
    current_section = None

    # 章节标题模式
    chapter_pattern = re.compile(r'^第[一二三四五六七八九十\d]+章\s*(.*)$')
    section_pattern = re.compile(r'^第[一二三四五六七八九十\d]+节\s*(.*)$')

    print(f"开始提取文档内容...")
    print(f"总段落数: {len(doc.paragraphs)}")

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # 检测章节标题
        chapter_match = chapter_pattern.match(text)
        if chapter_match:
            # 保存上一章节
            if current_chapter and current_section:
                current_chapter['sections'].append(current_section)
                current_section = None
            if current_chapter:
                chapters.append(current_chapter)

            # 开始新章节
            chapter_title = chapter_match.group(1).strip() if chapter_match.group(1) else text
            chapter_num = len(chapters) + 1

            current_chapter = {
                'chapter_id': f'ch{chapter_num:02d}',
                'chapter_name': text,
                'chapter_title': chapter_title,
                'sections': []
            }
            current_section = None
            print(f"发现章节: {text}")
            continue

        # 检测节标题
        section_match = section_pattern.match(text)
        if section_match and current_chapter:
            # 保存上一节
            if current_section:
                current_chapter['sections'].append(current_section)

            # 开始新节
            section_title = section_match.group(1).strip() if section_match.group(1) else text
            section_num = len(current_chapter['sections']) + 1

            current_section = {
                'section_id': f"{current_chapter['chapter_id']}_s{section_num:02d}",
                'section_name': text,
                'section_title': section_title,
                'text': ''
            }
            print(f"  发现节: {text}")
            continue

        # 普通段落内容
        if current_chapter:
            if current_section:
                # 添加到当前节
                current_section['text'] += text + '\n'
            else:
                # 如果没有节，创建默认节
                if not current_chapter['sections']:
                    current_section = {
                        'section_id': f"{current_chapter['chapter_id']}_s01",
                        'section_name': '正文',
                        'section_title': '正文',
                        'text': text + '\n'
                    }

    # 保存最后的章节和节
    if current_chapter and current_section:
        current_chapter['sections'].append(current_section)
    if current_chapter:
        chapters.append(current_chapter)

    print(f"\n提取完成！共 {len(chapters)} 个章节")

    return {
        'book_name': book_name,
        'book_id': book_id,
        'total_chapters': len(chapters),
        'chapters': chapters
    }


def extract_docx_to_json(docx_path: str, book_name: str, book_id: str, output_dir: Path) -> bool:
    """
    从Word文档提取内容并保存为JSON

    Args:
        docx_path: Word文档路径
        book_name: 书名
        book_id: 书籍ID
        output_dir: 输出目录

    Returns:
        是否成功
    """
    try:
        print(f"\n{'='*80}")
        print(f"处理文档: {book_name}")
        print(f"文件路径: {docx_path}")
        print(f"{'='*80}")

        doc = Document(docx_path)
        result = extract_chapter_structure(doc, book_name, book_id)

        if result['total_chapters'] == 0:
            print(f"警告: 未提取到任何章节")
            return False

        output_path = output_dir / f"{book_id}_chapters.json"
        output_dir.mkdir(parents=True, exist_ok=True)

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        print(f"\n成功保存到: {output_path}")
        print(f"章节数: {result['total_chapters']}")

        total_sections = sum(len(ch['sections']) for ch in result['chapters'])
        print(f"总节数: {total_sections}")

        return True

    except Exception as e:
        print(f"错误: 提取失败 - {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """主函数：处理三本教材"""
    base_dir = Path(__file__).parent.parent.parent.parent.parent / "文件"
    output_dir = Path(__file__).parent.parent / "data"

    textbooks = [
        {
            'file': '习近平新时代中国特色社会主义思想概论2023版.docx',
            'name': '习近平新时代中国特色社会主义思想概论（2023）',
            'id': 'xi_thought_2023'
        },
        {
            'file': '马克思主义基本原理 (2023年版)【马工程重点教材】.docx',
            'name': '马克思主义基本原理（2023年版）',
            'id': 'marxism_2023'
        },
        {
            'file': '思想道德与法治 (2023年版).docx',
            'name': '思想道德与法治（2023年版）',
            'id': 'morality_law_2023'
        }
    ]

    print(f"\n{'='*80}")
    print("开始提取三本教材内容")
    print(f"{'='*80}")
    print(f"源文件目录: {base_dir}")
    print(f"输出目录: {output_dir}")
    print(f"{'='*80}\n")

    success_count = 0
    for textbook in textbooks:
        docx_path = base_dir / textbook['file']
        if not docx_path.exists():
            print(f"警告: 文件不存在 - {docx_path}")
            continue

        if extract_docx_to_json(
            str(docx_path),
            textbook['name'],
            textbook['id'],
            output_dir
        ):
            success_count += 1

    print(f"\n{'='*80}")
    print(f"提取完成！成功: {success_count}/{len(textbooks)}")
    print(f"{'='*80}\n")


if __name__ == "__main__":
    main()
