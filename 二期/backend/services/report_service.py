# -*- coding: utf-8 -*-

"""
实践报告生成服务 —— 一键生成 Word 文档报告
"""

import json
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


PRACTICE_TYPE_MAP = {
    "writing": "写作设计类实践",
    "presentation": "宣传表达类实践",
    "visit": "参观研学类实践",
    "performance": "表演体验类实践",
    "interaction": "交往行动类实践",
    "production": "生产改造类实践",
    "free": "自由申请类实践",
}


def generate_report(submission) -> BytesIO:
    """
    根据 PracticeSubmission ORM 对象，生成一份 Word 报告并返回 BytesIO。
    报告格式参考: 133、表演体验类实践.docx 模板
    """
    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ---- 标题 ----
    title_text = "习近平新时代中国特色社会主义思想概论实践成果"
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title_text)
    run.font.size = Pt(16)
    run.bold = True
    run.font.name = "宋体"

    # ---- 基本信息表格 ----
    practice_type_label = PRACTICE_TYPE_MAP.get(
        submission.practice_type.value if hasattr(submission.practice_type, 'value') else str(submission.practice_type),
        str(submission.practice_type)
    )

    info_rows = [
        ("实践类别", practice_type_label),
        ("实践题目", submission.title or ""),
        ("成果形式", submission.result_form or ""),
        ("班级姓名学号", submission.class_name_id or ""),
        ("公众号展示", _showcase_label(submission.showcase_preference)),
        ("任课教师", submission.instructor_name or ""),
        ("提交日期", submission.completion_date.strftime("%Y年%m月%d日") if submission.completion_date else ""),
    ]

    table = doc.add_table(rows=len(info_rows), cols=2, style="Table Grid")
    for i, (label, value) in enumerate(info_rows):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        cell_label.text = label
        cell_value.text = value
        # 设置字体
        for cell in (cell_label, cell_value):
            for paragraph in cell.paragraphs:
                for r in paragraph.runs:
                    r.font.size = Pt(12)
                    r.font.name = "宋体"
        # 标签列宽度
        cell_label.width = Cm(3)

    doc.add_paragraph()  # 空行

    # ---- 正文内容 ----
    content_data = _parse_content(submission.content)

    # 如果 plan 关联的 tasks 有标题，尝试按步骤组织
    if content_data:
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = heading.add_run("实践过程与成果")
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = "宋体"

        for step_idx, step_texts in enumerate(content_data):
            if isinstance(step_texts, list):
                for text in step_texts:
                    if text and text.strip():
                        p = doc.add_paragraph()
                        r = p.add_run(text.strip())
                        r.font.size = Pt(12)
                        r.font.name = "宋体"
                        p.paragraph_format.line_spacing = Pt(20)
                        p.paragraph_format.first_line_indent = Cm(0.74)
            elif isinstance(step_texts, str) and step_texts.strip():
                p = doc.add_paragraph()
                r = p.add_run(step_texts.strip())
                r.font.size = Pt(12)
                r.font.name = "宋体"
                p.paragraph_format.line_spacing = Pt(20)
                p.paragraph_format.first_line_indent = Cm(0.74)

    # ---- 学生建议 ----
    if submission.reflection and submission.reflection.strip():
        doc.add_paragraph()
        heading = doc.add_paragraph()
        run = heading.add_run("学生建议")
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = "宋体"

        p = doc.add_paragraph()
        r = p.add_run(submission.reflection.strip())
        r.font.size = Pt(12)
        r.font.name = "宋体"
        p.paragraph_format.line_spacing = Pt(20)
        p.paragraph_format.first_line_indent = Cm(0.74)

    # ---- 输出 ----
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _showcase_label(pref: str) -> str:
    mapping = {
        "none": "不展示",
        "original": "原样展示",
        "anonymous": "匿名展示",
    }
    return mapping.get(pref or "original", pref or "")


def _parse_content(content_str: str):
    """解析 JSON 格式的 content（taskSubmissions 序列化结果）"""
    if not content_str:
        return []
    try:
        return json.loads(content_str)
    except (json.JSONDecodeError, TypeError):
        # 如果不是 JSON，当作纯文本
        return [[content_str]]
